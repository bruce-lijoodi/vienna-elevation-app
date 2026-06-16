"""
Generates the project technical documentation as a Word document.
Run with: py -3.14 generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
    p.paragraph_format.left_indent = Inches(0.4)

def spacer():
    doc.add_paragraph()

# ── Title page ───────────────────────────────────────────────────────────────

title = doc.add_heading("Vienna Elevation Router", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Technical Component Documentation")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

course = doc.add_paragraph("TU Wien — Location Based Services")
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course.runs[0].font.italic = True

doc.add_page_break()

# ── 1. Overview ──────────────────────────────────────────────────────────────

h1("1. Project Overview")
body(
    "The Vienna Elevation Router is a web application that calculates elevation-aware walking "
    "routes between two points in Vienna. Instead of finding only the shortest path, the app "
    "offers three route profiles — flattest, balanced, and steepest — so users can choose "
    "routes suited for accessibility, general walking, or fitness training respectively. "
    "The app runs in any browser without installation and is deployed publicly on Railway."
)
spacer()

# ── 2. Graph Builder ─────────────────────────────────────────────────────────

h1("2. Street Network Graph (graph_builder.py)")

h2("2.1 Downloading the Street Network")
body(
    "The app uses the OSMnx library to download Vienna's walkable street network directly "
    "from OpenStreetMap. The network is fetched as an 8 km radius circle centred on "
    "Stephansdom (48.2082°N, 16.3738°E), restricted to pedestrian-accessible paths only "
    "(network_type='walk'). The result is a directed multigraph with approximately 140,000 "
    "nodes (street intersections and path endpoints) and 300,000 edges (street segments)."
)
code("G = ox.graph_from_point((48.2082, 16.3738), dist=8000, network_type='walk', simplify=True)")
body(
    "The graph is serialised to disk using Python's pickle module as "
    "vienna_walk_graph_enriched.pkl (64 MB) so it only needs to be built once. "
    "On every subsequent server start, the file is loaded into RAM in a few seconds."
)
spacer()

h2("2.2 Elevation Data — SRTM CGIAR-CSI v4.1")
body(
    "Elevation values are sourced from the Shuttle Radar Topography Mission (SRTM) dataset, "
    "specifically the CGIAR-CSI version 4.1 post-processed tile srtm_40_03.tif. "
    "This tile covers central Europe at 3 arc-second (approximately 90 metre) horizontal "
    "resolution. Elevation is stored in metres above sea level as a raster grid."
)
body(
    "The rasterio library is used to read the GeoTIFF. The entire elevation grid is loaded "
    "into memory once, then for each graph node its latitude/longitude is converted to a "
    "row/column index in the raster and the elevation value is read directly."
)
code(
    "with rasterio.open('srtm_40_03.tif') as src:\n"
    "    elevation_data = src.read(1)          # load full grid once\n"
    "    for node_id, data in G.nodes(data=True):\n"
    "        row, col = src.index(data['x'], data['y'])   # lon, lat → pixel\n"
    "        G.nodes[node_id]['elevation'] = float(elevation_data[row, col])"
)
spacer()

h2("2.3 Edge Grade Calculation")
body(
    "After elevation is attached to every node, the steepness (grade) of each street "
    "segment is calculated and stored on the edge itself. Grade is defined as:"
)
code("grade = (elevation_end - elevation_start) / segment_length_metres")
body(
    "A positive grade means the segment goes uphill; negative means downhill. "
    "The absolute grade is also stored for use in filtering. Short segments with "
    "negligible elevation change are assigned grade = 0 to suppress SRTM raster noise."
)
spacer()

# ── 3. Routing Algorithm ─────────────────────────────────────────────────────

h1("3. Routing Algorithm (router.py)")

h2("3.1 Dijkstra's Algorithm with Custom Cost Functions")
body(
    "Route finding is handled by NetworkX's shortest_path function, which implements "
    "Dijkstra's algorithm. The key insight is that by changing the cost function — the "
    "function that assigns a 'price' to each edge — the same algorithm produces "
    "fundamentally different routes."
)
body(
    "The algorithm starts at the origin node and explores outward, always expanding the "
    "cheapest unvisited node, until it reaches the destination. The route with the lowest "
    "total cost is returned."
)
spacer()

h2("3.2 The Three Route Profiles")
body(
    "Three separate cost functions are defined, each expressing a different priority:"
)

bullet(
    "Flattest — uphill segments are made very expensive by adding a large penalty "
    "proportional to the elevation gain. The algorithm strongly prefers flat paths "
    "and will take long detours to avoid hills."
)
code("cost = length + 5000 × max(0, grade × length)")

bullet(
    "Balanced — pure shortest distance with no elevation consideration. "
    "This is a standard Dijkstra shortest-path."
)
code("cost = length")

bullet(
    "Steepest — uphill segments are made cheap (rewarded). The algorithm seeks out "
    "climbing sections. A minimum cost of 1.0 prevents negative weights."
)
code("cost = max(1.0, length − 5000 × max(0, grade × length))")

body(
    "The multiplier 5000 means that 1 metre of uphill climb is treated as equivalent to "
    "5000 metres of flat walking in terms of cost. This large value is needed because "
    "grade values are small decimals; without it the elevation penalty would be negligible "
    "compared to the edge length."
)
spacer()

h2("3.3 Nearest Node Snapping")
body(
    "Users specify locations by clicking the map or typing an address. These coordinates "
    "must be mapped to a node in the graph before routing can begin. The nearest_node "
    "function scans all nodes and returns the one with the smallest Euclidean distance "
    "in latitude/longitude space (sufficient precision for Vienna's scale)."
)
spacer()

h2("3.4 Path Statistics")
body(
    "Once a path (list of node IDs) is found, the _path_stats function walks through "
    "it to compute:"
)
bullet("Total distance in metres (summing edge lengths)")
bullet("Total elevation gain (sum of all uphill steps greater than 0.5 m)")
bullet("Total elevation loss (sum of all downhill steps greater than 0.5 m)")
bullet("Coordinate list [[lon, lat], ...] for drawing on the map")
bullet("Elevation profile: a list of {distance, elevation} points for the chart")
body(
    "The 0.5 m threshold filters out SRTM noise — very small elevation fluctuations "
    "between adjacent nodes that are artefacts of the 90 m raster resolution rather "
    "than real terrain."
)
spacer()

# ── 4. Loop Routes ────────────────────────────────────────────────────────────

h1("4. Circular / Loop Routes")

h2("4.1 Local Loop Algorithm (router.py)")
body(
    "When the user selects 'Take a Lap', the app generates circular routes that start "
    "and end at the same origin point. The local algorithm creates three loops by picking "
    "three compass directions (0°, 120°, 240°) and projecting a midpoint at half the "
    "target distance in each direction."
)
body(
    "For each direction, the algorithm routes from origin to the midpoint and back, "
    "then stitches the two sub-paths together (removing the duplicate junction node). "
    "The three resulting loops are sorted by elevation gain and labelled "
    "flattest / balanced / steepest."
)
code(
    "dlat_per_m = 1 / 111000\n"
    "dlon_per_m = 1 / (111000 × cos(origin_lat))\n"
    "mid_lat = origin_lat + (distance/2) × dlat_per_m × cos(angle)\n"
    "mid_lon = origin_lon + (distance/2) × dlon_per_m × sin(angle)"
)
spacer()

h2("4.2 GraphHopper Loop Algorithm (router_gh.py)")
body(
    "When GraphHopper is selected, the app uses GraphHopper's built-in round_trip "
    "routing algorithm. Three requests are made with different random seeds (0, 42, 123), "
    "which causes the algorithm to explore different parts of the surrounding terrain. "
    "The three routes are again sorted by elevation gain and labelled."
)
code(
    '"algorithm": "round_trip",\n'
    '"round_trip.distance": distance_metres,\n'
    '"round_trip.seed": 0   # or 42, or 123'
)
spacer()

# ── 5. GraphHopper Router ─────────────────────────────────────────────────────

h1("5. GraphHopper Routing Engine (router_gh.py)")
body(
    "GraphHopper is a commercial routing API that uses its own map data and a "
    "high-resolution Mapterhorn DEM (~30 m resolution) for elevation. It serves as "
    "an alternative to the local SRTM-based engine, offering finer elevation detail."
)
body(
    "For point-to-point routes, the app calls GraphHopper's alternative_route algorithm "
    "requesting up to 3 path variants. These are returned sorted by elevation gain. "
    "The algorithm finds genuinely different paths (not just minor detours) by penalising "
    "route similarity."
)
code(
    '"algorithm": "alternative_route",\n'
    '"alternative_route.max_paths": 3,\n'
    '"alternative_route.max_weight_factor": 2.0,\n'
    '"alternative_route.max_share_factor": 0.8'
)
body(
    "GraphHopper returns coordinates as [longitude, latitude, elevation] triples. "
    "The elevation value in the third coordinate is used to build the elevation profile, "
    "replacing the need to look up SRTM values separately."
)
body(
    "The API key is stored as an environment variable (GH_API_KEY) and never hard-coded. "
    "If the key is not set, the local OSMnx engine is used automatically."
)
spacer()

# ── 6. Navigation ─────────────────────────────────────────────────────────────

h1("6. Real-Time Navigation")

h2("6.1 GPS Tracking")
body(
    "Navigation uses the browser's Geolocation API (navigator.geolocation.watchPosition) "
    "to receive continuous position updates. High-accuracy GPS mode is intentionally "
    "disabled (enableHighAccuracy: false) because it causes indefinite hangs on laptops "
    "and desktops that lack a GPS chip. Network-based location (Wi-Fi/cell triangulation) "
    "is accurate enough for pedestrian routing."
)
spacer()

h2("6.2 Remaining Distance and ETA")
body(
    "On each position update, the app finds the nearest point on the route using the "
    "Haversine formula, which computes great-circle distance between two lat/lon points:"
)
code(
    "a = sin²(Δlat/2) + cos(lat1) × cos(lat2) × sin²(Δlon/2)\n"
    "distance = 2R × atan2(√a, √(1−a))   where R = 6,371,000 m"
)
body(
    "All remaining edge lengths from that index to the end of the route are summed to "
    "give the remaining distance. ETA is calculated by dividing remaining distance by "
    "the current GPS speed; if GPS speed is unavailable or below 0.3 m/s (stationary), "
    "a walking speed of 5 km/h (1.389 m/s) is assumed."
)
spacer()

h2("6.3 Route Split Visualisation")
body(
    "As the user walks, the route on the map is split into two segments at the nearest "
    "point. The completed portion is drawn as a grey dashed line; the remaining portion "
    "retains the profile colour (green / orange / red) at full opacity. "
    "The original three route lines are hidden during navigation."
)
spacer()

h2("6.4 Directional Arrow Marker")
body(
    "The user's live position is shown as a blue triangle SVG icon that rotates to face "
    "the direction of travel using the heading value from the Geolocation API. "
    "When the device is stationary (heading = null), the last known heading is preserved "
    "so the arrow does not snap back to north."
)
code('html: `<svg><polygon ... style="transform:rotate(${heading}deg)"/></svg>`')
spacer()

# ── 7. Elevation Profile Chart ────────────────────────────────────────────────

h1("7. Elevation Profile Chart")
body(
    "The elevation profile is rendered using Chart.js as a filled line chart. "
    "The X axis shows cumulative distance along the route in kilometres; "
    "the Y axis shows elevation in metres above sea level. "
    "The chart updates instantly when the user clicks a different route card, "
    "and uses the same colour as the route line on the map (green/orange/red)."
)
body(
    "Data comes from the elevation_profile field in the API response — a list of "
    "{distance, elevation} objects computed during path traversal in router.py. "
    "Individual data points are hidden (pointRadius: 0) and a slight tension (0.3) "
    "is applied to smooth the line."
)
spacer()

# ── 8. Address Search ─────────────────────────────────────────────────────────

h1("8. Address Search & Geocoding")
body(
    "Address search is powered by the Nominatim geocoding API (OpenStreetMap's free "
    "geocoder). As the user types, a debounced request fires after 400 ms of inactivity, "
    "querying Nominatim with the entered text plus ', Vienna' to bias results locally. "
    "Up to five suggestions are shown in a dropdown; selecting one sets the coordinates "
    "and places a marker on the map."
)
body(
    "Reverse geocoding is not required — the app works with raw lat/lon coordinates "
    "internally, and addresses are only used for display in the input field."
)
spacer()

# ── 9. Backend API ────────────────────────────────────────────────────────────

h1("9. Backend API (main.py)")

h2("9.1 FastAPI Framework")
body(
    "The backend is built with FastAPI, a modern Python web framework that automatically "
    "validates request and response data using Pydantic models, and generates interactive "
    "API documentation at /docs."
)
spacer()

h2("9.2 Graph Loading on Startup")
body(
    "When the server starts, a startup event loads the enriched graph pickle file into "
    "a global variable G. All subsequent route requests use this in-memory graph, "
    "avoiding repeated disk reads. If the pickle file does not exist, the graph is "
    "built from scratch (takes ~5 minutes on first run)."
)
spacer()

h2("9.3 POST /routes Endpoint")
body(
    "The main endpoint accepts a JSON body with:"
)
bullet("origin_lat, origin_lon — starting point")
bullet("dest_lat, dest_lon — destination (optional for loop mode)")
bullet("router — 'local' or 'graphhopper'")
bullet("mode — 'point_to_point' or 'loop'")
bullet("loop_distance_km — target loop length (default 5.0 km)")
body(
    "The endpoint validates that coordinates fall within Vienna's bounding box, "
    "selects the appropriate routing function, and returns a RouteResponse containing "
    "three RouteResult objects, each with distance, elevation gain/loss, coordinates, "
    "and elevation profile."
)
spacer()

# ── 10. Deployment ────────────────────────────────────────────────────────────

h1("10. Deployment")
body(
    "The app is deployed on Railway, a cloud platform that hosts the FastAPI server "
    "and serves the frontend HTML file from the same origin. A Procfile tells Railway "
    "how to start the server:"
)
code("web: uvicorn backend.main:app --host 0.0.0.0 --port $PORT")
body(
    "The 64 MB enriched graph file is committed directly to the git repository (not "
    "Git LFS) so Railway can access it during deployment. The GraphHopper API key is "
    "set as a Railway environment variable (GH_API_KEY) and never stored in code."
)

# ── Save ──────────────────────────────────────────────────────────────────────

out_path = r"d:\TUWIENNOTES\LOCATION BASED SERVICES\vienna-elevation-app\Vienna_Elevation_Router_Technical_Documentation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
